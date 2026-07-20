"""Local file intake. Extraction happens on this machine, before any policy runs.

Written outside Codex after the Codex quota was exhausted; see README
"How this was built".

Nothing here touches the network. A file is read from disk, its text is
extracted locally, and only then does the sanitizer and the preflight loop
see it. A document that cannot be parsed is refused rather than partially
imported, because a half-extracted document would be sanitized against text
the operator never saw.
"""

from __future__ import annotations

import logging
from pathlib import Path


class UnsupportedDocumentError(ValueError):
    """Raised when a file extension has no local extractor."""


class DocumentExtractionError(ValueError):
    """Raised when a supported file cannot be read as text."""


SUPPORTED_SUFFIXES = (".txt", ".md", ".pdf", ".docx")

# Formats deliberately out of scope for now. Named so the error message can
# say "not yet" rather than "unsupported", which is the honest distinction.
PLANNED_SUFFIXES = {
    ".xlsx": "spreadsheets need per-sheet and per-cell policy handling",
    ".xls": "spreadsheets need per-sheet and per-cell policy handling",
    ".pptx": "slide decks need per-slide extraction",
    ".png": "images need a vision model to read text",
    ".jpg": "images need a vision model to read text",
    ".jpeg": "images need a vision model to read text",
}


def ensure_supported(filename: str | Path) -> str:
    """Return the normalized suffix, or raise if nothing can extract it.

    Shared by the CLI and the upload endpoint so both refuse a format with the
    same message.
    """
    suffix = Path(filename).suffix.lower()
    if suffix in PLANNED_SUFFIXES:
        raise UnsupportedDocumentError(
            f"{suffix} is not supported yet: {PLANNED_SUFFIXES[suffix]}. "
            f"Supported today: {', '.join(SUPPORTED_SUFFIXES)}"
        )
    if suffix not in SUPPORTED_SUFFIXES:
        raise UnsupportedDocumentError(
            f"no local extractor for {suffix or 'a file with no extension'}. "
            f"Supported today: {', '.join(SUPPORTED_SUFFIXES)}"
        )
    return suffix


def extract_text(path: str | Path) -> str:
    """Return the plain text of a local document.

    Raises UnsupportedDocumentError for formats with no extractor and
    DocumentExtractionError when a supported format cannot be parsed.
    """
    file_path = Path(path).expanduser()
    if not file_path.is_file():
        raise DocumentExtractionError(f"not a file: {file_path}")

    suffix = ensure_supported(file_path)
    if suffix in {".txt", ".md"}:
        text = _read_plain_text(file_path)
    elif suffix == ".pdf":
        text = _read_pdf(file_path)
    else:
        text = _read_docx(file_path)

    if not text.strip():
        raise DocumentExtractionError(
            f"{file_path.name} produced no text. A scanned PDF needs OCR, "
            "which this tool does not do."
        )
    return text


def _read_plain_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError as error:
        raise DocumentExtractionError(f"{path.name} is not valid UTF-8 text") from error


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as error:
        raise DocumentExtractionError(
            "PDF support needs the optional dependency: pip install -e '.[files]'"
        ) from error

    # pypdf logs malformed-PDF warnings to stdout. Real client PDFs are often
    # malformed in harmless ways, and the operator does not need the noise.
    logging.getLogger("pypdf").setLevel(logging.ERROR)
    try:
        reader = PdfReader(str(path))
        if reader.is_encrypted:
            raise DocumentExtractionError(f"{path.name} is password protected")
        pages = [page.extract_text() or "" for page in reader.pages]
    except DocumentExtractionError:
        raise
    except Exception as error:
        raise DocumentExtractionError(f"could not read {path.name}: {error}") from error
    return "\n\n".join(page.strip() for page in pages if page.strip())


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as error:
        raise DocumentExtractionError(
            "Word support needs the optional dependency: pip install -e '.[files]'"
        ) from error

    try:
        document = Document(str(path))
    except Exception as error:
        raise DocumentExtractionError(f"could not read {path.name}: {error}") from error

    blocks = [paragraph.text.strip() for paragraph in document.paragraphs]
    # Table cells routinely hold the confidential values in consulting decks,
    # so they are extracted rather than skipped.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n".join(block for block in blocks if block)
